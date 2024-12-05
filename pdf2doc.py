"""
读取pdf文件，并将其中的文字、图片提取放到Word文档中
安装包：
pip install PyMuPDF python-docx
"""
import fitz
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# 复用pdf2md.py中的一些基础函数
def is_page_number(text, bbox, page_height, page_width):
    """判断文本是否为页码"""
    text = text.strip()
    if text.isdigit():
        y_pos = bbox[1]
        x_center = (bbox[0] + bbox[2]) / 2
        
        is_bottom = y_pos > page_height * 0.8
        is_center = page_width * 0.3 < x_center < page_width * 0.7
        
        return is_bottom and is_center
    return False

def get_text_features(span):
    """分析文本的特征"""
    return {
        'size': span['size'],
        'font': span['font'],
        'color': span['color'],
        'flags': span['flags'],
        'text': span['text']
    }

def analyze_text_style(spans):
    """分析一组文本的式特征"""
    features = []
    for span in spans:
        features.append(get_text_features(span))
    
    # 计算主要字体大小
    sizes = [f['size'] for f in features]
    main_size = max(set(sizes), key=sizes.count)
    
    # 改进粗体判断逻辑
    # 1. 检查字体名称中是否包含 Bold 关键字
    # 2. 检查flags中的粗体标志
    # 3. 要求主要文本都是粗体，而不是仅有部分文本是粗体
    total_text_length = sum(len(f['text']) for f in features)
    bold_text_length = sum(
        len(f['text']) 
        for f in features 
        if (f['flags'] & 2) or  # 粗体标志
           ('bold' in f['font'].lower()) or  # 字体名包含bold
           ('heavy' in f['font'].lower())    # 字体名包含heavy
    )
    
    # 只有当超过70%的文本是粗体时才认为整体是粗体
    is_bold = (bold_text_length / total_text_length) > 0.7 if total_text_length > 0 else False
    
    # 判断是否包含斜体
    is_italic = any(f['flags'] & 1 for f in features)  # 1 表示斜体
    
    return {
        'main_size': main_size,
        'is_bold': is_bold,
        'is_italic': is_italic,
        'text': ''.join(f['text'] for f in features)
    }

def is_title(block, font_sizes):
    """判断文本块是否为标题"""
    if not block.get("lines"):
        return False
        
    # 分析第一行的样式特征
    first_line = block["lines"][0]
    text_style = analyze_text_style(first_line["spans"])
    text = text_style['text'].strip()
    
    # 检查是否只有一行
    if len(block["lines"]) > 1:
        return False
    
    # 检查是否以数字开头，后面跟着标题文本
    if re.match(r'^\d+[、.]?\s*\S+', text) and len(text) <= 10:
        print(f"数字开头的标题-1：{text}")
        return True
    
    # 检查是否全部是粗体且以特定词开头
    if text_style['is_bold'] and any(text.startswith(word) for word in [
        '症状', '病原菌', '防治方法', '特征', '形态', '发生规律', '防治措施'
    ]):
        print(f"粗体标题：{text}")
        return True
    return False

def get_title_level(block, font_sizes):
    """根据文本特征确定标题级别"""
    if not block.get("lines"):
        return 6
        
    text_style = analyze_text_style(block["lines"][0]["spans"])
    current_size = text_style['main_size']
    text = text_style['text'].strip()
    
    # 按字体大小排序
    sorted_sizes = sorted(set(font_sizes), reverse=True)
    
    try:
        base_level = sorted_sizes.index(current_size) + 1
        
        # 根据其他特征调整级别
        if text_style['is_bold']:
            base_level = max(1, base_level - 1)  # 粗体提升一级
        if text_style['text'].strip().startswith(('第', '章')):
            base_level = 1  # 章节标题设为一级
        elif text_style['text'].strip().startswith(('节', '小节')):
            base_level = 2  # 小节标题设为二级
        return min(base_level, 6)
    except ValueError:
        return 6
    
def build_image_groups(images, page):
    """构建图片组，只处理横向相邻的图片"""
    if not images:
        return []
    
    def get_image_info(img_index, img):
        bbox = page.get_image_bbox(img)
        if not bbox:
            return None
        return {
            'index': img_index,
            'bbox': bbox,
            'center_x': (bbox[0] + bbox[2]) / 2,
            'center_y': (bbox[1] + bbox[3]) / 2,
            'width': bbox[2] - bbox[0],
            'height': bbox[3] - bbox[1]
        }
    
    # 获取所有图片的信息
    image_infos = []
    for idx, img in enumerate(images):
        info = get_image_info(idx, img)
        if info:
            image_infos.append(info)
    
    if not image_infos:
        return []
    
    # 按y坐标排序，处理每一行的图片
    image_infos.sort(key=lambda x: x['center_y'])
    
    # 分行处理
    rows = []
    current_row = [image_infos[0]]
    y_tolerance = 20  # 垂直方向的容差值
    
    for i in range(1, len(image_infos)):
        current_img = image_infos[i]
        last_img = current_row[-1]
        
        vertical_diff = abs(current_img['center_y'] - last_img['center_y'])
        if vertical_diff <= y_tolerance:
            current_row.append(current_img)
        else:
            current_row.sort(key=lambda x: x['bbox'][0])
            rows.append(current_row)
            current_row = [current_img]
    
    # 添加最后一行
    if current_row:
        current_row.sort(key=lambda x: x['bbox'][0])
        rows.append(current_row)
    
    # 处理每一行中的图片，将水平相邻的图片组合在一起
    groups = []
    horizontal_gap_threshold = 30  # 水平间距阈值
    
    for row in rows:
        current_group = [row[0]]
        
        for i in range(1, len(row)):
            current_img = row[i]
            last_img = current_group[-1]
            
            horizontal_gap = current_img['bbox'][0] - last_img['bbox'][2]
            
            if horizontal_gap < horizontal_gap_threshold:
                current_group.append(current_img)
            else:
                groups.append(current_group)
                current_group = [current_img]
        
        if current_group:
            groups.append(current_group)
    
    return groups

def find_image_caption(page, image_bbox, blocks):
    """查找图片下方的说明文字"""
    image_bottom = image_bbox[3]
    image_center_x = (image_bbox[0] + image_bbox[2]) / 2
    image_width = image_bbox[2] - image_bbox[0]
    
    caption_distance_threshold = 30
    horizontal_tolerance = image_width * 0.5
    
    potential_captions = []
    for block in blocks:
        if block.get("lines"):
            block_bbox = block["bbox"]
            block_top = block_bbox[1]
            block_center_x = (block_bbox[0] + block_bbox[2]) / 2
            
            vertical_distance = block_top - image_bottom
            horizontal_distance = abs(block_center_x - image_center_x)
            
            if (0 <= vertical_distance <= caption_distance_threshold and 
                horizontal_distance <= horizontal_tolerance):
                
                text = " ".join(span["text"] for line in block["lines"] 
                              for span in line["spans"]).strip()
                
                is_caption = False
                if any(keyword in text.lower() for keyword in ['图', 'fig', 'figure', '表', 'table']):
                    is_caption = True
                elif (len(text) < 100 and 
                      vertical_distance < 20 and 
                      block_bbox[2] - block_bbox[0] <= image_width * 1.2):
                    is_caption = True
                elif re.match(r'^\d+[.-]', text.strip()):
                    is_caption = True
                
                if is_caption:
                    potential_captions.append((text, block_bbox, vertical_distance))
    
    if potential_captions:
        return min(potential_captions, key=lambda x: x[2])[0:2]
    return None, None

def extract_images(page, output_dir):
    """提取页面中的图片及其说明文字"""
    image_list = []
    blocks = page.get_text("dict")["blocks"]
    
    page_width = page.rect.width
    page_height = page.rect.height
    
    images = list(page.get_images(full=True))
    image_groups = build_image_groups(images, page)
    
    for group_index, group in enumerate(image_groups):
        try:
            min_x = float('inf')
            min_y = float('inf')
            max_x = float('-inf')
            max_y = float('-inf')
            
            for img in group:
                bbox = img['bbox']
                min_x = min(min_x, bbox[0])
                min_y = min(min_y, bbox[1])
                max_x = max(max_x, bbox[2])
                max_y = max(max_y, bbox[3])
            
            original_min_x = min_x
            original_max_x = max_x
            
            min_x -= 2
            min_y -= 2
            max_x += 2
            max_y += 2
            
            left_margin = 5
            for block in blocks:
                if block.get("lines"):
                    block_bbox = block["bbox"]
                    if (block_bbox[3] > min_y and block_bbox[1] < max_y):
                        if block_bbox[2] < min_x:
                            min_x = max(min_x, block_bbox[2] + left_margin)
            
            caption_text, caption_bbox = find_image_caption(page, [min_x, min_y, max_x, max_y], blocks)
            
            if caption_bbox:
                max_y = caption_bbox[3] + 2
                min_x = original_min_x - 2
                max_x = original_max_x + 2
            
            clip_bbox = [
                max(0, min_x),
                max(0, min_y),
                min(page_width, max_x),
                min(page_height, max_y)
            ]
            
            if clip_bbox[2] <= clip_bbox[0] or clip_bbox[3] <= clip_bbox[1]:
                print(f"警告：页面 {page.number + 1} 的图片组 {group_index} 截取区域无效，跳过")
                continue
            
            width = clip_bbox[2] - clip_bbox[0]
            height = clip_bbox[3] - clip_bbox[1]
            if width < 1 or height < 1 or width > page_width or height > page_height:
                print(f"警告：页面 {page.number + 1} 的图片组 {group_index} 尺寸无效，跳过")
                continue
            
            try:
                # 修改图片保存部分
                try:
                    # 提高分辨率以保持图片质量
                    zoom = 3.0  # 增加分辨率
                    matrix = fitz.Matrix(zoom, zoom)
                    # 添加错误检查
                    if (clip_bbox[2] - clip_bbox[0]) < 1 or (clip_bbox[3] - clip_bbox[1]) < 1:
                        print(f"警告：页面 {page.number + 1} 的图片组 {group_index} 区域太小，跳过")
                        continue
                        
                    try:
                        # 使用更高质量的设置获取图片
                        pix = page.get_pixmap(matrix=matrix, clip=clip_bbox, alpha=False)
                    except Exception as e:
                        print(f"警告：获取页面 {page.number + 1} 的图片组 {group_index} 像素数据失败: {str(e)}")
                        continue
                        
                    if not pix or pix.width < 1 or pix.height < 1:
                        print(f"警告：页面 {page.number + 1} 的图片组 {group_index} 像素数据无效")
                        continue
                        
                    image_filename = f"image_{page.number + 1}_{group_index}_group.png"
                    image_path = os.path.join(output_dir, "images", image_filename)
                    os.makedirs(os.path.dirname(image_path), exist_ok=True)
                    
                    try:
                        # 直接使用 PyMuPDF 保存高质量图片
                        pix.save(image_path, output="png")
                        
                        # 使用 PIL 优化图片大小但保持较高质量
                        try:
                            from PIL import Image
                            with Image.open(image_path) as img:
                                # 使用更高的质量设置
                                img.save(image_path, format='PNG', optimize=True, quality=95)
                        except Exception as e:
                            print(f"警告：优化图片 {image_path} 失败: {str(e)}")
                            # 即使优化失败也继续使用原图
                        
                        image_list.append({
                            "path": os.path.join("images", image_filename),
                            "bbox": clip_bbox,
                            "caption": caption_text
                        })
                        
                    except Exception as e:
                        print(f"警告：保存图片 {image_path} 失败: {str(e)}")
                        continue
                        
                except Exception as e:
                    print(f"警告：处理页面 {page.number + 1} 的图片组 {group_index} 时出错: {str(e)}")
                    continue
                    
            except Exception as e:
                print(f"警告：提取页面 {page.number + 1} 的图片时出错: {str(e)}")
                
        except Exception as e:
            print(f"警告：处理页面 {page.number + 1} 的图片组 {group_index} 时出错: {str(e)}")
            continue
    
    return image_list

def clean_text_for_docx(text):
    """清理文本，移除不兼容的字符"""
    # 移除控制字符，但保留换行和制表符
    text = ''.join(char for char in text if char >= ' ' or char in ['\n', '\t'])
    # 替换 NULL 字节
    text = text.replace('\x00', '')
    # 合并连续的换行符
    import re
    # 先将所有连续的换行符（包括可能的空白字符）合并为单个换行符
    text = re.sub(r'\r?\n\s*\r?\n+', '\r\n', text)
    # 将剩余的单个换行符替换为段落符
    text = re.sub(r'\r?\n', '\r\n', text)
    return text

def process_text_block(block, font_sizes, page_height, page_width, image_captions=None):
    """处理单个文本块"""
    if not block.get("lines"):
        return "", False
    
    # 检查是否为图片说明
    if image_captions:
        block_bbox = block["bbox"]
        for caption_bbox in image_captions:
            if (block_bbox[0] >= caption_bbox[0] - 1 and 
                block_bbox[2] <= caption_bbox[2] + 1 and 
                block_bbox[1] >= caption_bbox[1] - 1 and 
                block_bbox[3] <= caption_bbox[3] + 1):
                return "", False

    text = ""
    is_bold = True
    
    for line in block["lines"]:
        line_text = ""
        line_is_bold = True
        
        for span in line["spans"]:
            span_text = span["text"].strip()
            if span_text:
                # 清理文本
                span_text = clean_text_for_docx(span_text)
                line_text += span_text
                if not ((span['flags'] & 2) or 
                       ('bold' in span['font'].lower()) or 
                       ('heavy' in span['font'].lower())):
                    line_is_bold = False
        
        if line_text and not is_page_number(line_text, block["bbox"], page_height, page_width):
            if text and (text.endswith('-') or text.endswith('－')):
                text = text[:-1]
            text += line_text
            is_bold = is_bold and line_is_bold
    
    return clean_text_for_docx(text.strip()), is_bold

def pdf_to_doc(pdf_path, output_dir):
    """将PDF文件转换为Word文档"""
    os.makedirs(output_dir, exist_ok=True)
    
    doc = fitz.open(pdf_path)
    document = Document()
    
    # 设置文档基本样式
    style = document.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
    doc_path = os.path.join(output_dir, f"{pdf_name}.docx")
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        
        page_height = page.rect.height
        page_width = page.rect.width
        
        # 处理图片
        images = extract_images(page, output_dir)
        image_captions = []
        for img in images:
            if img.get("caption") and img.get("bbox"):
                image_captions.append(img["bbox"])
        
        # 收集字体大小
        font_sizes = []
        blocks = page.get_text("dict")["blocks"]
        for block in blocks:
            if block.get("lines"):
                for line in block["lines"]:
                    for span in line["spans"]:
                        font_sizes.append(span["size"])
        
        # 处理文本，传入图片说明位置列表
        current_text = []
        non_title_text = []  # 存储非标题文本
        
        for block in blocks:
            if block.get("lines"):
                # 判断是否为标题
                    if is_title(block, font_sizes):
                        # 处理累积的非标题文本
                        if non_title_text:
                            merged_text = ""
                            for text in non_title_text:
                                if any(merged_text.endswith(end) for end in ['。', '！', '？', '.', '!', '?', '：', ':', ';', '；']):
                                    merged_text += '\r\n' + text
                                else:
                                    merged_text += text
                            current_text.append(merged_text + '\r\n\r\n')
                            non_title_text = []
                        
                        # 处理标题
                        level = get_title_level(block, font_sizes)
                        text, _ = process_text_block(block, font_sizes, page_height, page_width, image_captions)
                        if text:
                            print(f"识别到标题 {level}: {text}")
                            current_text.append(f"{'#' * level} {text}\r\n\r\n")
                    else:
                        # 检查是否为加粗文本（子标题）
                        text, is_bold = process_text_block(block, font_sizes, page_height, page_width, image_captions)
                        if text:
                            if is_bold:
                                # 处理之前累积的非标题文本
                                if non_title_text:
                                    merged_text = ""
                                    for prev_text in non_title_text:
                                        if any(merged_text.endswith(end) for end in ['。', '！', '？', '.', '!', '?', '：', ':', ';', '；']):
                                            merged_text += '\r\n' + prev_text
                                        else:
                                            merged_text += prev_text
                                    current_text.append(merged_text + '\r\n\r\n')
                                    non_title_text = []
                                # 添加加粗文作为子标题
                                current_text.append(f"**{text}**\r\n\r\n")
                            else:
                                non_title_text.append(text)
        # 处理最后剩余的非标题文本
        if non_title_text:
            merged_text = ""
            for text in non_title_text:
                if any(merged_text.endswith(end) for end in ['。', '！', '？', '.', '!', '?', '：', ':', ';', '；']):
                    merged_text += '\r\n' + text
                else:
                    merged_text += text
            current_text.append(merged_text + '\r\n\r\n')
        # 修改写入文本的部分
        # 将:
        # document.write("".join(current_text))
        # 改为:
        for text in current_text:
            # 处理标题
            if text.startswith('#'):
                level = text.count('#')
                title_text = text[level:].strip()
                document.add_heading(title_text, level=level)
            # 处理加粗文本
            elif text.startswith('**') and text.endswith('**\r\n\r\n'):
                p = document.add_paragraph()
                p.add_run(text[2:-4]).bold = True
            # 处理普通文本
            else:
                document.add_paragraph(text)
        
        # 添加图片
        for img in images:
            try:
                # 根据页面宽度调整图片尺寸
                page_width = document.sections[0].page_width
                max_width = page_width * 0.8  # 图片宽度设为页面宽度的80%
                document.add_picture(
                    os.path.join(output_dir, img['path']), 
                    width=min(Inches(6), max_width)
                )
            except Exception as e:
                print(f"警告：添加图片时出错: {str(e)}")
        
        # 添加分页符
        if page_num < len(doc) - 1:
            document.add_page_break()
    
    # 保存文档
    document.save(doc_path)
    doc.close()
    return doc_path

def batch_convert_pdfs(input_dir, output_base_dir):
    """批量转换文件夹中的PDF文件为Word文档"""
    os.makedirs(output_base_dir, exist_ok=True)
    
    pdf_files = [f for f in os.listdir(input_dir) if f.lower().endswith('.pdf')]
    
    if not pdf_files:
        print(f"在 {input_dir} 中没有找到PDF文件")
        return
    
    for pdf_file in pdf_files:
        pdf_path = os.path.join(input_dir, pdf_file)
        pdf_name = os.path.splitext(pdf_file)[0]
        output_dir = os.path.join(output_base_dir, pdf_name)
        
        try:
            doc_path = pdf_to_doc(pdf_path, output_dir)
            print(f"成功转换 {pdf_file} -> {doc_path}")
        except Exception as e:
            print(f"转换 {pdf_file} 时出错: {str(e)}")

if __name__ == '__main__':
    input_dir = "D:\\pdfs"  # PDF文件所在文件夹
    output_dir = "D:\\outputs"  # 输出的基础目录
    batch_convert_pdfs(input_dir, output_dir)