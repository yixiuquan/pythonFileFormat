import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import logging
from datetime import datetime
import shutil

# 配置日志
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler(f'pdf_extract_{datetime.now().strftime("%Y%m%d_%H%M%S")}.log')
    ]
)
logger = logging.getLogger(__name__)

def create_output_dirs(pdf_dir, pdf_name):
    """创建输出目录结构
    Args:
        pdf_dir: PDF文件所在目录
        pdf_name: PDF文件名（不含扩展名）
    Returns:
        tuple: (pdf输出目录, 图片输出目录)
    """
    # 创建PDF同名目录
    pdf_output_dir = os.path.join(pdf_dir, pdf_name)
    if os.path.exists(pdf_output_dir):
        # 如果目录已存在，先删除
        shutil.rmtree(pdf_output_dir)
    os.makedirs(pdf_output_dir)
    
    # 创建imgs子目录
    imgs_dir = os.path.join(pdf_output_dir, 'imgs')
    os.makedirs(imgs_dir)
    
    return pdf_output_dir, imgs_dir

def is_yellow_square(annot):
    """���断是否为黄色矩形标注
    Args:
        annot: PDF注释对象
    Returns:
        bool: 是否为黄色矩形标注
    """
    try:
        # 检查是否是矩形注释（type[0] == 4 表示Square）
        if annot.type[0] != 4:
            return False
            
        # 检查颜色信息
        colors = annot.colors
        if not colors:
            return False
            
        # 检查填充颜色是否为黄色
        fill_color = colors.get('fill')
        if not fill_color:
            return False
            
        r, g, b = fill_color
        # 判断是否为黄色（R和G接近1，B接近0）
        return (r > 0.9 and g > 0.9 and b < 0.1)
        
    except Exception as e:
        logger.debug(f"检查注释类型时出错: {str(e)}")
        return False

def extract_highlights_from_pdf(pdf_path, output_dir):
    """从PDF文件中提取高亮区域"""
    try:
        # 打开原始文档以获取注释信息
        doc = fitz.open(pdf_path)
        highlight_blocks = []
        
        pdf_name = os.path.splitext(os.path.basename(pdf_path))[0]
        # 创建输出目录
        pdf_output_dir, imgs_dir = create_output_dirs(output_dir, pdf_name)

        # 创建一个没有任何注释的新文档
        clean_doc = fitz.open()
        for page in doc:
            # 复制页面但不复制注释
            clean_doc.new_page(width=page.rect.width, 
                             height=page.rect.height, 
                             pno=-1).show_pdf_page(
                                 page.rect,
                                 doc,
                                 page.number
                             )

        # 使用原始文档获取注释位置，使用清洁文档进行截图
        for page_num in range(len(doc)):
            page = doc[page_num]
            clean_page = clean_doc[page_num]
            logger.debug(f"处理第 {page_num + 1} 页")
            
            # 获取页面上的所有注释
            annotations = list(page.annots())
            if annotations:
                logger.debug(f"第 {page_num + 1} 页找到 {len(annotations)} 个注释")
                
                for annot_num, annot in enumerate(annotations):
                    if is_yellow_square(annot):
                        # 获取高亮区域的矩形框
                        rect = annot.rect
                        
                        # 保存截图
                        img_filename = f"page_{page_num + 1}_highlight_{len(highlight_blocks) + 1}.png"
                        img_path = os.path.join(imgs_dir, img_filename)
                        
                        # 扩大截取区域以确保完整捕获内容
                        clip_rect = fitz.Rect(rect)
                        clip_rect.x0 -= 3
                        clip_rect.x1 += 3
                        clip_rect.y0 -= 3
                        clip_rect.y1 += 3
                        
                        # 使用清洁页面进行截图
                        pix = clean_page.get_pixmap(matrix=fitz.Matrix(3, 3), clip=clip_rect)
                        pix.save(img_path)
                        
                        highlight_blocks.append({
                            'page': page_num + 1,
                            'image_path': img_path
                        })
                        logger.info(f"在第 {page_num + 1} 页找到黄色矩形区域")

        doc.close()
        clean_doc.close()
        
        if not highlight_blocks:
            logger.warning("未找到任何黄色矩形区域，请确认PDF中是否包含黄色矩形标注")
            # 如果没有找到高亮区域，删除创建的目录
            shutil.rmtree(pdf_output_dir)
            return [], None
        return highlight_blocks, pdf_output_dir

    except Exception as e:
        logger.error(f"处理PDF文件 {pdf_path} 时出错: {str(e)}", exc_info=True)
        return [], None

def save_to_word(highlight_blocks, output_dir, pdf_name):
    """将高亮区域的截图保存到Word文档"""
    try:
        doc = Document()
        
        # 直接添加所有截图
        for block in highlight_blocks:
            if os.path.exists(block['image_path']):
                doc.add_picture(block['image_path'], width=Inches(6))
                doc.add_paragraph()  # 添加空行分隔图片

        # 保存文档到PDF同名目录
        word_path = os.path.join(output_dir, f"{pdf_name}_标注提取.docx")
        doc.save(word_path)
        logger.info(f"已保存Word文档: {word_path}")
        return True

    except Exception as e:
        logger.error(f"保存Word文档时出错: {str(e)}", exc_info=True)
        return False

def process_pdf_directory(pdf_dir):
    """处理目录下的所有PDF文件"""
    try:
        if not os.path.exists(pdf_dir):
            logger.error(f"目录不存在: {pdf_dir}")
            return

        # 获取所有PDF文件
        pdf_files = [f for f in os.listdir(pdf_dir) if f.lower().endswith('.pdf')]
        
        if not pdf_files:
            logger.warning(f"在 {pdf_dir} 中未找到PDF文件")
            return

        # 处理每个PDF文件
        for pdf_file in pdf_files:
            pdf_path = os.path.join(pdf_dir, pdf_file)
            pdf_name = os.path.splitext(pdf_file)[0]
            logger.info(f"开始处理: {pdf_file}")

            # 提取高亮区域
            highlight_blocks, pdf_output_dir = extract_highlights_from_pdf(pdf_path, pdf_dir)
            
            if highlight_blocks:
                if save_to_word(highlight_blocks, pdf_output_dir, pdf_name):
                    logger.info(f"成功处理文件: {pdf_file}")
                else:
                    logger.error(f"保存Word文档失败: {pdf_file}")
            else:
                logger.warning(f"未在 {pdf_file} 中找到黄色标注区域")

        logger.info("所有文件处理完成")

    except Exception as e:
        logger.error(f"处理目录时出错: {str(e)}", exc_info=True)

if __name__ == "__main__":
    PDF_DIR = "D:/pdfs"  # 固定的PDF目录路径
    logger.info(f"开始处理目录: {PDF_DIR}")
    process_pdf_directory(PDF_DIR)
    logger.info("处理完成")